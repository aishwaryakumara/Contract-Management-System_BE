"""
Contract Extraction Service using NER
=====================================

Uses spaCy's pre-trained NER model (en_core_web_sm) to extract entities from contracts.

Supported formats: PDF, DOCX, DOC

Model: en_core_web_sm v3.8.0
- Trained on: OntoNotes 5.0 (news, web, broadcast, telephone)
- Size: 12.8 MB
- Entities: PERSON, ORG, GPE, DATE, MONEY, CARDINAL, etc.
- Architecture: CNN (Convolutional Neural Network)
- Accuracy: ~85% on benchmark datasets

How It Works:
1. Extract text from PDF (PyPDF2) or Word (python-docx)
2. Run spaCy NER pipeline:
   - Tokenization (split into words)
   - POS tagging (noun, verb, etc.)
   - Dependency parsing (grammar structure)
   - Named Entity Recognition (identify entities)
3. Extract specific entities we need for contracts
4. Apply heuristics for accuracy
"""

import re
import os
import spacy
from typing import Dict, Optional, List
from datetime import datetime
from dateutil import parser as date_parser
from PyPDF2 import PdfReader
from docx import Document  # For Word documents


class ContractExtractionService:
    """Extract contract metadata from PDFs and Word documents using NER"""
    
    def __init__(self):
        """
        Initialize the NER model
        
        Loads: en_core_web_sm
        Components loaded:
        - tok2vec: Token-to-vector embeddings
        - tagger: Part-of-speech tagger
        - parser: Dependency parser
        - ner: Named Entity Recognizer (this is what we use!)
        - attribute_ruler: Set custom attributes
        - lemmatizer: Word lemmatization
        """
        print("🔄 Loading spaCy NER model...")
        self.nlp = spacy.load("en_core_web_sm")
        print("✅ Model loaded successfully!")
        
        # Entity types we care about
        self.target_entities = ['ORG', 'DATE', 'MONEY', 'PERSON']
    
    def extract_from_pdf(self, file_path: str) -> Dict:
        """
        Main extraction method (supports PDF and Word documents)
        
        Args:
            file_path: Path to document file (PDF, DOCX, DOC)
            
        Returns:
            Dictionary with extracted fields:
            {
                'client_name': str,
                'contract_name': str,
                'contract_type': str,
                'start_date': str (YYYY-MM-DD),
                'end_date': str (YYYY-MM-DD),
                'value': float,
                'description': str
            }
        """
        print(f"\n📄 Extracting from: {file_path}")
        
        # Step 1: Detect file type and extract text
        file_type = self._detect_file_type(file_path)
        print(f"📋 File type: {file_type.upper()}")
        
        if file_type == 'pdf':
            text = self._extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            text = self._extract_text_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Use PDF or DOCX")
        
        print(f"📝 Extracted {len(text)} characters")
        
        # Step 2: Run NER on text
        doc = self.nlp(text[:5000])  # Limit to first 5000 chars for speed
        print(f"🤖 NER found {len(doc.ents)} entities")
        
        # Step 3: Extract specific fields
        result = {
            'client_name': self._extract_client_name(doc, text),
            'contract_name': self._extract_contract_name(text),
            'contract_type': self._extract_contract_type(text),
            'start_date': self._extract_start_date(doc),
            'end_date': self._extract_end_date(doc),
            'value': self._extract_value(doc),
            'description': self._extract_description(text)
        }
        
        print(f"✅ Extraction complete!")
        return result
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using PyPDF2
        
        How it works:
        - Opens PDF file
        - Iterates through each page
        - Extracts text using PDF text layer
        - Concatenates all pages
        """
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"❌ Error reading PDF: {str(e)}")
            return ""
        
        return text
    
    def _extract_text_from_word(self, file_path: str) -> str:
        """
        Extract text from Word document (.docx) using python-docx
        
        How it works:
        - Opens Word document
        - Iterates through paragraphs
        - Extracts text from each paragraph
        - Handles tables (extracts table content)
        - Concatenates all text
        
        Note: Only works with .docx (not old .doc format)
        """
        text = ""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                text += "\n"
            
        except Exception as e:
            print(f"❌ Error reading Word document: {str(e)}")
            return ""
        
        return text
    
    def _detect_file_type(self, file_path: str) -> str:
        """
        Detect file type based on extension
        
        Returns: 'pdf', 'docx', or 'unknown'
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return 'pdf'
        elif ext in ['.docx', '.doc']:
            return 'docx'
        else:
            return 'unknown'
    
    def _extract_client_name(self, doc, text: str) -> Optional[str]:
        """
        Extract client/company name using NER
        
        Strategy:
        1. Get all ORG entities from spaCy (organizations)
        2. Filter out common false positives
        3. Look for contract-specific patterns
        4. Return most likely client name
        
        spaCy ORG entity examples:
        - "Acme Corporation" → ORG
        - "Google Inc" → ORG
        - "XYZ Holdings LLC" → ORG
        """
        # Get all organizations detected by NER
        organizations = [ent.text for ent in doc.ents if ent.label_ == 'ORG']
        
        if organizations:
            # Filter out common false positives
            filtered_orgs = [
                org for org in organizations 
                if org.lower() not in ['agreement', 'contract', 'llc']
            ]
            
            if filtered_orgs:
                print(f"   🏢 Found organizations: {filtered_orgs[:3]}")
                # Return first valid organization (usually the client)
                return filtered_orgs[0] if len(filtered_orgs) >= 1 else filtered_orgs[0]
        
        # Fallback: Look for "Client:" pattern with regex
        pattern = r'(?:Client|Customer|Company):\s*([A-Z][A-Za-z\s&.,]+?)(?:\n|$)'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        
        return None
    
    def _extract_contract_name(self, text: str) -> Optional[str]:
        """
        Extract contract title/name
        
        Strategy:
        - Look for first heading with "AGREEMENT" or "CONTRACT"
        - Usually in first 500 characters
        - Must be reasonable length (10-100 chars)
        """
        lines = text.split('\n')
        
        # Look for lines with contract keywords
        for line in lines[:10]:
            line = line.strip()
            if any(kw in line.upper() for kw in ['AGREEMENT', 'CONTRACT']):
                if 10 < len(line) < 100:
                    print(f"   📋 Contract name: {line}")
                    return line
        
        return "Service Agreement"  # Default fallback
    
    def _extract_contract_type(self, text: str) -> Optional[str]:
        """
        Classify contract type using keyword matching
        
        Types supported:
        - service: Service agreements, consulting
        - purchase: Purchase orders, sales
        - license: License agreements, software
        - nda: Non-disclosure, confidentiality
        
        Method: Count keyword occurrences and pick highest
        """
        text_lower = text.lower()
        
        type_keywords = {
            'service': ['service agreement', 'services', 'consulting', 'professional services'],
            'purchase': ['purchase order', 'purchase agreement', 'sales agreement', 'buy'],
            'license': ['license agreement', 'licensing', 'software license', 'intellectual property'],
            'nda': ['non-disclosure', 'nda', 'confidentiality agreement', 'confidential']
        }
        
        # Score each type
        scores = {}
        for contract_type, keywords in type_keywords.items():
            scores[contract_type] = sum(1 for kw in keywords if kw in text_lower)
        
        # Return type with highest score
        if max(scores.values()) > 0:
            best_type = max(scores, key=scores.get)
            print(f"   📑 Contract type: {best_type} (score: {scores[best_type]})")
            return best_type
        
        return 'service'  # Default
    
    def _parse_valid_dates(self, doc) -> List[str]:
        """
        Extract and validate all dates from document
        
        Returns only valid, parseable dates (filters out false positives like zip codes)
        """
        date_entities = [ent.text for ent in doc.ents if ent.label_ == 'DATE']
        valid_dates = []
        
        for date_str in date_entities:
            try:
                # Try to parse the date
                date_obj = date_parser.parse(date_str, fuzzy=False)
                
                # Validate it's a reasonable date
                if 2000 <= date_obj.year <= 2100:  # Reasonable year range
                    formatted = date_obj.strftime('%Y-%m-%d')
                    valid_dates.append(formatted)
            except:
                # Not a valid date (e.g., "MA 02115")
                continue
        
        return valid_dates
    
    def _extract_start_date(self, doc) -> Optional[str]:
        """
        Extract start date using spaCy DATE entity
        
        spaCy DATE entity examples:
        - "January 15, 2024" → DATE
        - "15th of January 2024" → DATE
        - "01/15/2024" → DATE
        - "next month" → DATE (relative)
        
        Strategy:
        1. Get all DATE entities from NER
        2. Filter out false positives (zip codes, etc.)
        3. Return first valid date
        """
        valid_dates = self._parse_valid_dates(doc)
        
        if valid_dates:
            print(f"   📅 Valid dates found: {valid_dates}")
            print(f"   ✅ Start date: {valid_dates[0]}")
            return valid_dates[0]
        
        return None
    
    def _extract_end_date(self, doc) -> Optional[str]:
        """
        Extract end date (usually second DATE entity)
        
        Strategy:
        1. Get all DATE entities
        2. Filter out false positives
        3. Return second valid date (first is start date)
        """
        valid_dates = self._parse_valid_dates(doc)
        
        if len(valid_dates) >= 2:
            print(f"   ✅ End date: {valid_dates[1]}")
            return valid_dates[1]
        
        return None
    
    def _extract_value(self, doc) -> Optional[float]:
        """
        Extract contract value using spaCy MONEY entity
        
        spaCy MONEY entity examples:
        - "$50,000" → MONEY
        - "fifty thousand dollars" → MONEY
        - "$50K" → MONEY
        - "USD 50000" → MONEY
        
        Strategy:
        1. Get all MONEY entities from NER
        2. Extract numeric value
        3. Handle K/M suffixes (50K = 50000)
        """
        money_entities = [ent.text for ent in doc.ents if ent.label_ == 'MONEY']
        
        if money_entities:
            print(f"   💰 Found money: {money_entities[:3]}")
            
            for money_str in money_entities:
                try:
                    # Remove currency symbols and commas
                    clean = re.sub(r'[$,\s]', '', money_str)
                    
                    # Handle K/M suffixes
                    if 'K' in clean.upper():
                        clean = clean.upper().replace('K', '000')
                    if 'M' in clean.upper():
                        clean = clean.upper().replace('M', '000000')
                    
                    # Extract number
                    match = re.search(r'(\d+(?:\.\d+)?)', clean)
                    if match:
                        value = float(match.group(1))
                        print(f"   ✅ Value: ${value:,.2f}")
                        return value
                except:
                    continue
        
        return None
    
    def _extract_description(self, text: str) -> Optional[str]:
        """
        Extract brief description (first paragraph or summary)
        
        Strategy:
        - Take first 200 characters after title
        - Remove extra whitespace
        - Return as description
        """
        # Get first paragraph (after title)
        paragraphs = text.split('\n\n')
        if len(paragraphs) >= 2:
            desc = paragraphs[1].strip()
            # Limit to 200 chars
            if len(desc) > 200:
                desc = desc[:200] + '...'
            return desc
        
        return None
    
    def get_model_info(self) -> Dict:
        """
        Get information about the loaded model
        
        Useful for debugging and showing what's under the hood
        """
        return {
            'name': self.nlp.meta['name'],
            'version': self.nlp.meta['version'],
            'language': self.nlp.meta['lang'],
            'pipeline': self.nlp.pipe_names,
            'entities': list(self.nlp.get_pipe('ner').labels)
        }
